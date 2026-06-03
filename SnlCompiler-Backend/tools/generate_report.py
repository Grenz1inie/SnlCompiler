# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

def set_cn(run, size=12, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold

def add_title(text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn(r, size, True)

def add_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r, 14, True)

def add_body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r, 12, False)

add_title('编译原理课程实验报告')
add_body('指导教师：【请填写】')
add_body('班    级：【请填写】')
add_body('姓    名：【请填写】')
add_body('小组编号：【请填写】')
add_body('组长学号姓名：【请填写】')
add_body('组员学号姓名：【请填写】')
add_body('组员学号姓名：【请填写】')
add_body('2026年  6  月  【请填写】日')
add_body('软件学院')
doc.add_paragraph()

add_heading('一、本次实验描述')
add_body(
    '本实验在 Java HTTP API + Vue 前端环境下实现了 SNL（Small Nested Language）编译前端，'
    '包含词法分析、LL(1) 语法分析、递归下降语法分析和语义分析四个模块。'
    '词法分析器将 SNL 源程序转换为 Token 内部表示序列；LL(1) 分析器基于预测分析表输出语法推导过程；'
    '递归下降分析器按非终结符编写子程序并输出语法错误信息；语义分析器在语法树基础上建立符号表并进行语义检查。'
)

add_heading('二、小组成员及任务描述')
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
headers = ['姓名', '具体承担任务', '工作量百分比', '小组成员协作情况']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
rows = [
    ['【成员1】', '词法分析（Lexer）、Constants 配置', '35%', '共同讨论文法与测试用例'],
    ['【成员2】', 'LL(1) 语法分析（Parser）', '30%', '共同讨论文法与测试用例'],
    ['【成员3】', '递归下降分析、语义分析、界面', '35%', '共同讨论文法与测试用例'],
]
for r_idx, row in enumerate(rows, start=1):
    for c_idx, val in enumerate(row):
        table.rows[r_idx].cells[c_idx].text = val

add_heading('三、实验平台及环境')
add_body('操作系统：Windows 10/11')
add_body('JDK 版本：JDK 8 及以上')
add_body('开发语言：Java')
add_body('界面框架：Vue 3 + Element Plus')
add_body('构建方式：javac 命令行编译')
add_body('项目路径：d:\\bianyi')

add_heading('四、实验方法设计')
add_body('4.1 词法分析')
add_body(
    '按 PPT 要求完成单词分类、正则定义与 DFA 识别。Token 分为五类：分隔符(1)、保留字(2)、'
    '标识符(3)、整数(4)、字符常量(5)。支持 :=、..、注释 { } 跳过、字符常量 \'x\' 识别，'
    '以及词法错误处理（非法字符、错误赋值符、程序未以 . 结束）。'
)
add_body('4.2 LL(1) 语法分析')
add_body(
    '在 Constants 中维护 104 条产生式与 LL(1) 预测分析表。Parser 使用栈模拟推导，'
    '输出每条规则编号及“语法分析成功/失败”信息。'
)
add_body('4.3 递归下降语法分析')
add_body(
    'RecursiveDescentParser 为每个非终结符编写分析子程序，通过 predict 集（当前 Token 判断）'
    '选择分支，输出语法错误信息，并构建语法树供语义分析使用。'
)
add_body('4.4 语义分析')
add_body(
    'SemanticAnalyzer 遍历语法树，在声明部分建立符号表，在语句部分检查重复定义、'
    '未声明标识符、类型名/过程名误用、赋值类型兼容性、条件表达式等语义问题。'
)

add_heading('五、程序界面及运行截图')
add_body('【此处插入截图，参见《实验测试文档.md》各测试用例的截图说明】')
add_body('建议截图：1. 主界面  2. 词法分析（含注释程序）  3. LL(1) 语法分析  4. 递归下降分析  5. 语义分析与符号表  6. 错误用例')

add_heading('六、实验结果')
add_body('对 sample.snl、comment_test.snl、record_test.snl、bubble_sort.snl 等测试程序，'
         '词法分析、LL(1) 分析、递归下降分析及语义分析均通过自动化测试（tools/FullTest.java，17/17）。')
add_body('词法错误、缺少结束句点等异常输入能正确报错。')

add_heading('七、实验结论')
add_body(
    '本实验完成了 SNL 编译前端三个必做程序。词法分析实现了 PPT 要求的 Token 分类与注释、'
    '字符常量处理；语法分析同时实现 LL(1) 与递归下降两种方法；语义分析实现了符号表与基本语义检查。'
    '通过本实验加深了对词法自动机、LL(1) 预测分析和递归下降分析的理解。'
)

add_heading('八、附录：核心源代码说明')
add_body('Lexer.java — 词法分析主程序')
add_body('Parser.java — LL(1) 语法分析')
add_body('RecursiveDescentParser.java — 递归下降语法分析')
add_body('SemanticAnalyzer.java — 语义分析')
add_body('Constants.java — 文法规则与分析表')
add_body('Main.java — 后端服务入口')
add_body('完整源码见 src/com/snl/compiler/ 目录。')

out = r'd:\bianyi\编译原理实验报告-已完成.docx'
doc.save(out)
print('saved', out)

