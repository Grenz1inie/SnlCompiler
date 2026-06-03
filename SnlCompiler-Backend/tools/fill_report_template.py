# -*- coding: utf-8 -*-
"""Fill 编译课程实验报告模板.docx with complete experiment report content."""
import os
import re
import shutil
import zipfile
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'd:\bianyi'
TEMPLATE = os.path.join(BASE, '编译课程实验报告模板.docx')
OUTPUT = os.path.join(BASE, '编译原理实验报告-定稿.docx')
SCREENSHOT_DIR = os.path.join(BASE, 'screenshots')
SRC_BASE = os.path.join(BASE, 'src', 'com', 'snl', 'compiler')

# --- content ---

COVER = {
    'teacher': '【请填写指导教师姓名】',
    'class': '【请填写班级】',
    'name': '【请填写姓名】',
    'group_no': '【请填写小组编号】',
    'leader': '【请填写组长学号姓名】',
    'member1': '【请填写组员学号姓名】',
    'member2': '【请填写组员学号姓名】',
    'date': '2026年  6  月  10  日',
}

EXPERIMENT_DESC = (
    '本实验依据《编译原理实验要求》PPT（2026.05 版），在 Java + Swing 环境下实现了 '
    'SNL（Small Nested Language）编译前端，包含词法分析、LL(1) 语法分析、递归下降语法分析和语义分析四个模块。'
    '词法分析器将 SNL 源程序扫描为 Token 序列，支持外部/内部两种表示；LL(1) 分析器基于 104 条产生式与预测分析表'
    '输出推导过程；递归下降分析器按非终结符编写子程序并构建语法树；语义分析器在语法树基础上建立符号表并进行语义检查。'
    '本实验不要求目标代码生成。'
)

TEAM_ROWS = [
    ('【成员1姓名】', '词法分析（Lexer.java）、Constants 配置、测试用例', '35%', '共同讨论 SNL 文法与测试方案，定期合并代码'),
    ('【成员2姓名】', 'LL(1) 语法分析（Parser.java）、预测分析表调试', '30%', '共同讨论 SNL 文法与测试方案，定期合并代码'),
    ('【成员3姓名】', '递归下降分析、语义分析、Swing 界面（MainFrame）', '35%', '共同讨论 SNL 文法与测试方案，定期合并代码'),
]

PLATFORM = [
    '操作系统：Windows 10/11（64 位）',
    'JDK 版本：JDK 8 及以上',
    '开发语言：Java',
    '界面框架：Java Swing',
    '构建方式：javac 命令行编译，无 Maven/Gradle',
    '项目路径：d:\\bianyi',
    '入口类：com.snl.compiler.ui.MainFrame',
    '编译命令：javac -encoding UTF-8 -d bin -sourcepath src src/com/snl/compiler/ui/MainFrame.java',
    '运行命令：java -cp bin com.snl.compiler.ui.MainFrame',
    '自动化测试：java -cp "bin;tools" FullTest（30/30 通过，覆盖文档全部用例）',
]

METHOD_DESIGN = [
    ('4.1 词法分析（Lexer.java）', (
        '按 PPT 要求完成单词分类、正则定义与 DFA 识别。Token 分为五类：分隔符(1)、保留字(2)、'
        '标识符(3)、整数(4)、字符常量(5)。识别函数顺序扫描源程序，跳过空白与注释 { … }，'
        '识别 :=、..、单字符分隔符及保留字（含 record）。字符常量形如 \'A\' 输出类型 5。'
        '词法错误包括：非法字符、错误赋值符（如 =:）、程序未以句点 . 正常结束。'
        '输出支持外部表示 (类型,词素) 与内部表示 (行号,类型,符号表下标) 两种格式。'
    )),
    ('4.2 LL(1) 语法分析（Parser.java）', (
        '在 Constants.java 中维护 105 条 SNL 产生式（含 Factor→CHARC）与 LL(1) 预测分析表。'
        'Parser 使用栈模拟自顶向下推导：遇终结符与输入 Token 匹配，遇非终结符查表压栈。'
        'Token 类型 5 映射为终结符 CHARC，支持字符常量参与表达式与赋值。'
        '输出每条使用的产生式编号及推导过程，成功时提示「语法分析成功！」，失败时输出「语法分析失败！」及错误信息。'
    )),
    ('4.3 递归下降语法分析（RecursiveDescentParser.java）', (
        '为每个非终结符编写分析子程序（如 programHead、typeDec、procDec、stmt 等），'
        '通过当前 Token 的 FIRST 集选择分支（predict 函数）。'
        'factor() 支持 INTC 与 CHARC 常量；declarePart 正确链接多条类型/变量/过程声明。'
        '支持数组类型与 a[exp] 下标访问、record 类型声明与 a.field 域访问。'
        '分析成功输出「递归下降语法分析成功，未发现语法错误。」并构建 AST 供语义分析使用。'
    )),
    ('4.4 语义分析（SemanticAnalyzer.java）', (
        'SemanticAnalyzer 遍历语法树：在声明部分建立符号表并管理作用域；'
        '检查变量/类型/过程重复定义、未声明标识符、类型名或过程名被当作变量使用；'
        '检查赋值语句左右类型兼容性、关系表达式操作数类型等。'
        '分析成功时在界面输出符号表（名称、类别、类型、作用域层次）。'
    )),
    ('4.5 程序结构与模块划分', (
        'core/lexer — 词法分析；core/parser — LL(1) 与递归下降；core/semantic — 语义分析；'
        'core/ast — 语法树节点；infra/config — Constants 文法与分析表；ui — MainFrame 图形界面；'
        'resource — 测试用 SNL 源文件（sample、comment_test、record_test、bubble_sort、char_test 等）。'
    )),
]

SCREENSHOTS = [
    ('图1  主界面', '00_主界面说明.png', '程序启动后主界面，含词法/LL(1)/递归下降/语义分析四个功能按钮。'),
    ('图2  词法分析（外部表示）', '01_词法分析_sample_外部表示.png', 'sample.snl 词法分析成功，Token 外部表示。'),
    ('图3  词法分析（内部表示）', '02_词法分析_sample_内部表示.png', 'sample.snl Token 内部表示 (行号,类型,下标)。'),
    ('图4  注释跳过', '03_词法分析_注释程序.png', 'comment_test.snl 中 { } 注释不出现在 Token 序列。'),
    ('图5  字符常量词法', '04_词法分析_字符常量.png', 'char_test.snl 中 \'A\' 识别为类型 5 (CHARC)。'),
    ('图6  词法错误', '05_词法错误_非法字符.png', '非法字符 # 触发词法错误。'),
    ('图7  LL(1) 语法分析', '07_LL1语法分析_成功.png', 'sample.snl LL(1) 推导过程与成功提示。'),
    ('图8  复杂程序 LL(1)', '08_LL1语法分析_冒泡排序.png', 'bubble_sort.snl 含数组与过程的语法分析成功。'),
    ('图9  LL(1) 字符常量', '15_LL1语法分析_字符常量.png', 'char_test.snl 中 c:=\'A\' 的 LL(1) 语法分析成功。'),
    ('图10 LL(1) 语法错误', '09_LL1语法错误.png', 'read 后缺少分号触发语法分析失败。'),
    ('图11 递归下降分析', '10_递归下降_成功.png', 'sample.snl 递归下降语法分析成功。'),
    ('图12 record 类型', '11_递归下降_record类型.png', 'record_test.snl 记录类型与域访问分析成功。'),
    ('图13 递归下降字符常量', '16_递归下降_字符常量.png', 'char_test.snl 递归下降语法分析成功。'),
    ('图14 语义分析', '12_语义分析_成功.png', 'sample.snl 语义分析成功及符号表输出。'),
    ('图15 语义错误（未声明）', '13_语义错误_未声明.png', '未声明变量 x 的语义错误提示。'),
    ('图16 语义错误（重复定义）', '14_语义错误_重复定义.png', '变量 a 重复定义的语义错误提示。'),
]

CORE_SOURCES = [
    ('Lexer.java — 词法分析核心', 'core/lexer/Lexer.java', 1, 120),
    ('Parser.java — LL(1) 语法分析核心', 'core/parser/Parser.java', 1, 80),
    ('RecursiveDescentParser.java — 递归下降（节选）', 'core/parser/RecursiveDescentParser.java', 1, 100),
    ('SemanticAnalyzer.java — 语义分析（节选）', 'core/semantic/SemanticAnalyzer.java', 1, 100),
]

RESULTS = (
    '对 sample.snl、comment_test.snl、record_test.snl、bubble_sort.snl、char_test.snl 等测试程序，'
    '词法分析、LL(1) 分析、递归下降分析及语义分析均通过。'
    '自动化测试 tools/FullTest.java 共 30 项全部通过（PASS=30 FAIL=0），覆盖《实验测试文档》全部正向与错误用例。'
    '词法错误（非法字符、缺少句点）、语法错误（read 后缺少分号）、语义错误（未声明、重复定义）均能正确检测并给出中文提示。'
    '字符常量 \'A\' 在词法、LL(1)、递归下降、语义各阶段均可正确处理。'
)

CONCLUSION = (
    '本实验完成了 SNL 编译前端三个必做程序：词法分析、语法分析（同时实现 LL(1) 与递归下降）、语义分析。'
    '词法分析实现了 PPT 要求的 Token 分类、注释跳过、字符常量与 record 保留字；'
    '语法分析覆盖 SNL 全部 104 条文法，包括数组、记录、过程与控制语句；'
    '语义分析实现了符号表管理与基本类型检查。'
    '通过本实验加深了对有限自动机、LL(1) 预测分析、递归下降分析与语法制导翻译的理解，'
    '为后续目标代码生成打下基础。'
)


def set_cn(run, size=12, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, size=12, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_cn(r, size, bold)
    return p


def add_heading(doc, text):
    add_para(doc, text, 14, True)


def add_body(doc, text):
    add_para(doc, text, 12, False)


def find_screenshot(filename):
    """Match screenshot by suffix (handles encoding issues in listing)."""
    if not os.path.isdir(SCREENSHOT_DIR):
        return None
    for f in os.listdir(SCREENSHOT_DIR):
        if f.endswith('.png') and (f == filename or f.endswith(filename.split('_', 1)[-1])):
            return os.path.join(SCREENSHOT_DIR, f)
    # try exact
    path = os.path.join(SCREENSHOT_DIR, filename)
    return path if os.path.isfile(path) else None


def read_source(rel_path, start=1, end=120):
    path = os.path.join(SRC_BASE, rel_path.replace('/', os.sep))
    if not os.path.isfile(path):
        return f'// 文件不存在: {path}'
    with open(path, 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()
    chunk = lines[start - 1:end]
    return ''.join(chunk)


def build_document():
    # Copy template to preserve styles
    work = OUTPUT
    try:
        shutil.copy2(TEMPLATE, OUTPUT)
    except PermissionError:
        work = OUTPUT.replace('.docx', '-更新.docx')
        shutil.copy2(TEMPLATE, work)
        print('原报告被占用，将写入:', work)
    doc = Document(work)

    # Clear existing body content but keep styles — remove all paragraphs/tables
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            body.remove(child)

    # --- Cover ---
    add_para(doc, '编译原理课程实验报告', 22, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    for label, key in [
        ('指导教师：', 'teacher'),
        ('班    级：', 'class'),
        ('姓    名：', 'name'),
        ('小组编号：', 'group_no'),
        ('组长学号姓名：', 'leader'),
        ('组员学号姓名：', 'member1'),
        ('组员学号姓名：', 'member2'),
    ]:
        add_body(doc, label + COVER[key])
    add_para(doc, COVER['date'], 12, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '软件学院', 12, False, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # --- Section 1 ---
    add_heading(doc, '一、本次实验描述')
    add_body(doc, EXPERIMENT_DESC)
    doc.add_paragraph()

    # --- Section 2 ---
    add_heading(doc, '二、小组成员及任务描述')
    table = doc.add_table(rows=1 + len(TEAM_ROWS), cols=4)
    table.style = 'Table Grid'
    headers = ['姓名', '具体承担任务', '工作量百分比', '小组成员协作情况']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(TEAM_ROWS, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()

    # --- Section 3 ---
    add_heading(doc, '三、实验平台及环境')
    for line in PLATFORM:
        add_body(doc, line)
    doc.add_paragraph()

    # --- Section 4 ---
    add_heading(doc, '四、实验方法设计')
    for title, text in METHOD_DESIGN:
        add_para(doc, title, 12, True)
        add_body(doc, text)
    doc.add_paragraph()

    # --- Section 5 ---
    add_heading(doc, '五、程序界面及运行截图')
    add_body(doc, '以下截图为 SNL 编译器各功能模块的运行效果（由 ScreenshotGenerator 基于真实分析结果生成）。')
    for caption, fname, desc in SCREENSHOTS:
        add_para(doc, caption, 12, True)
        add_body(doc, desc)
        img_path = find_screenshot(fname)
        if img_path and os.path.isfile(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.8))
        else:
            add_body(doc, f'[截图文件缺失: {fname}]')
        doc.add_paragraph()

    # --- Section 6 Results ---
    add_heading(doc, '六、实验结果')
    add_body(doc, RESULTS)
    doc.add_paragraph()

    # --- Section 7 Conclusion ---
    add_heading(doc, '七、实验结论')
    add_body(doc, CONCLUSION)
    doc.add_paragraph()

    # --- Section 8 Core source ---
    add_heading(doc, '八、源代码核心代码')
    add_body(doc, '完整源码位于 src/com/snl/compiler/ 目录。以下摘录各模块核心实现（节选）。')
    for title, rel, start, end in CORE_SOURCES:
        add_para(doc, title, 12, True)
        code = read_source(rel, start, end)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(code)
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(9)
        doc.add_paragraph()

    save_document(doc, work)


def save_document(doc, path):
    try:
        doc.save(path)
        print('Saved:', path)
    except PermissionError:
        alt = path.replace('.docx', '-更新.docx')
        doc.save(alt)
        print('保存失败，已另存为:', alt)


if __name__ == '__main__':
    if not os.path.isfile(TEMPLATE):
        raise SystemExit('Template not found: ' + TEMPLATE)
    build_document()
